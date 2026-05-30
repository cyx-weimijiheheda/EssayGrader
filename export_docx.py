# -*- coding: utf-8 -*-
"""
高中英语应用文批改工具 — DOCX 导出模块
"""

import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


FONT_LATIN = "Calibri"
FONT_EAST = "等线"
FONT_SIZE_TITLE = Pt(15)
FONT_SIZE_H1 = Pt(12)
FONT_SIZE_H2 = Pt(11)
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)

MARGIN = Cm(1.8)

HEADER_BG = "D9E8F7"
MODULE_BG = "F5F5F5"
POLISH_BG = "F0F5FB"
LEARNING_BG = "FFF7E0"


def _set_font(run, size=FONT_SIZE_BODY, bold=False, color=None, strike=False):
    """设置 run 的西文+东亚字体和样式"""
    run.font.name = FONT_LATIN
    run.font.size = size
    run.bold = bold
    run.font.strike = strike
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_EAST)
    rPr.insert(0, rFonts)


def _set_spacing(paragraph, before=0, after=2, line=None):
    """设置段落间距（紧凑）"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _add_heading(doc, text, level=1):
    """添加标题段落"""
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = FONT_SIZE_TITLE
        _set_spacing(p, before=0, after=6)
    elif level == 1:
        size = FONT_SIZE_H1
        _set_spacing(p, before=10, after=4)
    else:
        size = FONT_SIZE_H2
        _set_spacing(p, before=8, after=2)
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    return p


def _add_body(doc, text=""):
    """添加正文段落"""
    p = doc.add_paragraph()
    if text:
        text = text.strip()
        run = p.add_run(text)
        _set_font(run)
    _set_spacing(p, before=0, after=2, line=1.35)
    return p


def _apply_bordered_style(paragraph, bg_color: str = None):
    """给段落添加左边框和可选背景色"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "8")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "4A90D9")
    pBdr.append(left)
    pPr.append(pBdr)
    if bg_color:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), bg_color)
        pPr.append(shd)


def _add_bordered_box(doc, text: str, bg_color: str = None):
    """在浅色边框框中添加文本"""
    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=2, line=1.4)
    if text:
        run = p.add_run(text.strip())
        _set_font(run)
    _apply_bordered_style(p, bg_color)
    return p


_INLINE_ERROR_RE = re.compile(r"\[错误:(.+?)→(.+?)\|(.+?)\]")


# ==================== 各板块渲染函数 ====================

def _setup_student_section(doc, section, student_name, barcode_data, export_timestamp):
    """为每位学生设置独立的页眉"""
    if doc.sections[0] != section:
        section.top_margin = doc.sections[0].top_margin
        section.bottom_margin = doc.sections[0].bottom_margin
        section.left_margin = doc.sections[0].left_margin
        section.right_margin = doc.sections[0].right_margin

    header = section.header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p.clear()

    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    parts = []
    parts.append(f"姓名：{student_name}" if student_name else "姓名：未知")
    if barcode_data:
        parts.append(f"考号：{barcode_data}")
    parts.append(export_timestamp)

    text = " | ".join(parts)
    run = p.add_run(text)
    _set_font(run, size=FONT_SIZE_SMALL)
    _set_spacing(p, before=0, after=0)


def _add_essay_title_line(doc, essay_title):
    """在每位学生报告中添加作文题目"""
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=4)
    run = p.add_run(f"作文题目：{essay_title}")
    _set_font(run, size=FONT_SIZE_SMALL, bold=True, color=RGBColor(0x33, 0x33, 0x33))


def _add_score_table(doc, total, content, language, structure, fmt):
    """添加评分详情表格 — 紧凑双行表头+数值"""
    _add_heading(doc, "评分详情", level=2)

    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["总分", "内容要点", "语言质量", "篇章结构", "格式与语体"]
    values = [f"{total}/15", f"{content}/5", f"{language}/5", f"{structure}/3", f"{fmt}/2"]

    for i in range(5):
        cell_h = table.cell(0, i)
        cell_h.text = ""
        p_h = cell_h.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_h = p_h.add_run(headers[i])
        _set_font(r_h, size=FONT_SIZE_SMALL, bold=True)
        _set_spacing(p_h, before=1, after=1)
        tc = cell_h._element
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_BG)
        tcPr.append(shd)

        cell_v = table.cell(1, i)
        cell_v.text = ""
        p_v = cell_v.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_v = p_v.add_run(values[i])
        _set_font(r_v, size=FONT_SIZE_BODY, bold=(i == 0))
        if i == 0:
            r_v.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        _set_spacing(p_v, before=1, after=1)

    _set_spacing(doc.add_paragraph(), before=0, after=0)


def _add_error_fallback(doc, errors: list):
    """无修正版时的简化错误列表"""
    _add_heading(doc, "错误分析", level=2)
    if errors:
        text = "\n".join(f"{i}. {e}" for i, e in enumerate(errors, 1))
        _add_bordered_box(doc, text, bg_color=MODULE_BG)
    else:
        _add_body(doc, "（无）")


def _add_corrected_section(doc, corrected_text: str):
    """添加修正版板块，错误处红色+删除线，统一边框样式"""
    _add_heading(doc, "修正版（含错误标注）", level=2)

    if not corrected_text:
        _add_body(doc, "（未生成改错修正版）")
        return

    parts = _INLINE_ERROR_RE.split(corrected_text)

    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=2, line=1.45)
    for i, part in enumerate(parts):
        if part == "":
            continue
        rem = i % 4
        if rem == 0:
            run = p.add_run(part)
            _set_font(run)
        elif rem == 1:
            run = p.add_run(part)
            _set_font(run, color=RGBColor(0xCC, 0x00, 0x00), strike=True)
        elif rem == 2:
            run = p.add_run(" → " + part)
            _set_font(run, color=RGBColor(0x00, 0x88, 0x00))
        elif rem == 3:
            run = p.add_run("（" + part + "）")
            _set_font(run, size=FONT_SIZE_SMALL, color=RGBColor(0x99, 0x99, 0x99))
    _apply_bordered_style(p, bg_color=MODULE_BG)


def _add_comment_section(doc, comment: str):
    """添加评语板块，统一边框样式"""
    _add_heading(doc, "评语", level=2)
    if comment:
        _add_bordered_box(doc, comment, bg_color=MODULE_BG)
    else:
        _add_body(doc, "（未生成评语）")


def _add_polished_section(doc, polished_text: str):
    """添加精修升格范文板块"""
    _add_heading(doc, "精修升格范文", level=2)

    if not polished_text:
        _add_body(doc, "（未生成升格范文）")
        return

    _add_bordered_box(doc, polished_text, bg_color=POLISH_BG)


def _add_learning_section(doc, learning_text: str):
    """添加学习版范文板块（主题偏离时出现）"""
    _add_heading(doc, "学习版范文（直接回应题目要求）", level=2)

    if not learning_text:
        return

    _add_bordered_box(doc, learning_text, bg_color=LEARNING_BG)


# ==================== 主导出函数 ====================

def export_to_docx(all_results: list, essay_title: str, filepath: str):
    """导出批改结果为 DOCX 文件，每位学生一个独立节（含个性化页眉）"""
    doc = Document()

    # ---- 页面设置 ----
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    # ---- 默认样式 ----
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST)

    export_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # ---- 逐位学生输出 ----
    for idx, entry in enumerate(all_results):
        result = entry.get("result", {})

        total = result.get("total_score", 0)
        content = result.get("content_score", 0)
        language = result.get("language_score", 0)
        structure = result.get("structure_score", 0)
        fmt = result.get("format_score", 0)
        corrected = result.get("corrected_version", "")
        errors = result.get("errors", [])
        comment = result.get("comment", "")
        polished = result.get("polished_version", "")
        learning = result.get("learning_version", "")
        student_name = result.get("student_name", "")
        student_class = result.get("student_class", "")
        barcode_data = result.get("barcode_data", "")

        # ---- 节管理 ----
        if idx == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()

        _setup_student_section(doc, section, student_name, barcode_data, export_timestamp)

        # ---- 学生报告标题 ----
        display_name = student_name or f"第 {idx + 1} 篇"
        title = display_name
        if student_class:
            title += f" | {student_class}"
        title += f"（得分：{total}/15分）"
        _add_heading(doc, title, level=1)

        # ---- 作文题目 ----
        _add_essay_title_line(doc, essay_title)

        # 1. 评分详情
        _add_score_table(doc, total, content, language, structure, fmt)

        # 2. 评语
        if comment:
            _add_comment_section(doc, comment)

        # 3. 修正版（含错误标注）
        if corrected:
            _add_corrected_section(doc, corrected)
        elif errors:
            _add_error_fallback(doc, errors)

        # 4. 精修升格范文
        if polished:
            _add_polished_section(doc, polished)

        # 5. 学习版范文（主题偏离时）
        if learning:
            _add_learning_section(doc, learning)

    doc.save(filepath)
